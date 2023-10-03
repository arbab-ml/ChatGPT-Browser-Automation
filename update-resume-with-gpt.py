#To create new chrome user: 
# /Applications/Google\ Chrome.app/Contents/MacOS/Google\ Chrome --user-data-dir="/Users/muhammadarbabarshad/temporarychromeuser"
# /Applications/Google\ Chrome.app/Contents/MacOS/Google\ Chrome --user-data-dir="/Users/muhammadarbabarshad/temporarychromeuser-arbab-isu"
from docx import Document
from handler.chatgpt_selenium_automation import ChatGPTAutomation

def get_job_experiences(doc_path):
    doc = Document(doc_path)
    start_processing = False
    
    experiences = []
    current_experience = []
    for paragraph in doc.paragraphs:
        if "Relevant Experience" in paragraph.text:
            start_processing = True
            continue
        if start_processing:
            if not paragraph.text.strip() and current_experience:
                experiences.append(current_experience)
                current_experience = []
            else:
                current_experience.append(paragraph)

    if current_experience:
        experiences.append(current_experience)
    
    return experiences

def process_experience_with_chatgpt(experience_paragraphs, chrome_path, chrome_driver_path):
    with open('job-description.txt', 'r') as jd_file:
        job_description = jd_file.read()
    
    bullets_string = "\n".join([para.text for para in experience_paragraphs])
    
    base_prompt = """
    Adjust the experience in my resume as per the job description. Keep it the same length or less. Just optimize the keywords to get over ATS and do not add any fake experience. Do not add or imply experience or skills that I did not mention. Job roles and specific achievements should not be altered but merely rephrased for keyword optimization. Be creative in keywords but ensure that they are correct in the context of my experience. The output should contain only the revised bullet points and no other text. The number of output bullet points will be the same as input bullet points.
    """
    
    prompt = base_prompt + "\nJob Description:\n" + job_description + "\nCurrent Experience Bullet Points:\n" + bullets_string
    
    chatgpt = ChatGPTAutomation(chrome_path, chrome_driver_path)
    chatgpt.send_prompt_to_chatgpt(prompt)
    response = chatgpt.return_last_response()
    chatgpt.quit()
    
    return response.split("\n")

def update_document_with_processed_experience(doc_path, original_paragraphs, processed_experiences):
    doc = Document(doc_path)
    extract_experience_index=0
    for i, paragraph in enumerate(original_paragraphs):
        if not paragraph.text.strip():
            continue
        if i < len(processed_experiences):
            # Check if the paragraph has multiple runs
            if len(paragraph.runs) > 1:
                # Set the text of the first run to the entire processed experience
                paragraph.runs[0].text = processed_experiences[extract_experience_index]
                # Clear the text of subsequent runs
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.text = processed_experiences[extract_experience_index]

    doc.save("updated-resume.docx")

def main():
    doc_path = "resume.docx"
    chrome_driver_path = "//Users/muhammadarbabarshad/ChatGPT-Browser-Automation/chromedriver-mac-arm64/chromedriver"
    chrome_path = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
    
    all_experiences = get_job_experiences(doc_path)
    processed_bullets = process_experience_with_chatgpt(all_experiences[0], chrome_path, chrome_driver_path)
    update_document_with_processed_experience(doc_path, all_experiences[0], processed_bullets)
    print("Done")

if __name__ == "__main__":
    main()