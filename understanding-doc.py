from docx import Document
import json

def placeholder_function(bullet_json):
    """Capitalize the first word of each bullet."""
    bullet_list = json.loads(bullet_json)
    
    # Process each bullet to capitalize its first word
    processed_bullets = []
    for bullet in bullet_list:
        first_word, remaining_text = bullet.split(' ', 1)
        processed_bullet = first_word.capitalize() + ' ' + remaining_text
        processed_bullets.append(processed_bullet)
    
    return json.dumps(processed_bullets)

def update_bullet_points(doc_path):
    doc = Document(doc_path)
    
    # Flag to start collecting paragraphs under job experiences
    start_collecting = False
    bullet_points = []
    
    for paragraph in doc.paragraphs:
        # Check if we're in the 'Relevant Experience' section or other job sections
        if "Relevant Experience" in paragraph.text or "Additional Experience" in paragraph.text:
            start_collecting = True
            continue  # Skip the title
        
        if start_collecting and paragraph.text.strip() == "":
            start_collecting = False  # Reset flag when we encounter a blank line
        
        if start_collecting:
            bullet_points.append(paragraph.text)
    
    # Convert bullet_points to JSON and process them
    bullet_json = json.dumps(bullet_points)
    processed_bullet_json = placeholder_function(bullet_json)
    processed_bullets = json.loads(processed_bullet_json)
    
    # Replace original bullets with processed bullets
    idx = 0
    start_replacing = False
    for paragraph in doc.paragraphs:
        # Check if we're in the 'Relevant Experience' section or other job sections
        if "Relevant Experience" in paragraph.text or "Additional Experience" in paragraph.text:
            start_replacing = True
            continue
        
        if start_replacing and paragraph.text.strip() == "":
            start_replacing = False
        
        if start_replacing:
            paragraph.text = processed_bullets[idx]
            idx += 1
    
    # Save the updated document to a new file
    doc.save("updated_" + doc_path)

# Call the function
doc_path = "resume.docx"
update_bullet_points(doc_path)